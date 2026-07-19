import fitz  # PyMuPDF
import docx  # python-docx
import io
from typing import Union

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text content from PDF bytes using PyMuPDF."""
    text = ""
    # Open PDF document from memory
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from DOCX bytes using python-docx."""
    text_parts = []
    # Open Word document from memory
    doc = docx.Document(io.BytesIO(file_bytes))
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_parts.append(paragraph.text)
            
    # Extract text from tables (optional, but helpful for CVs with table layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text_parts.append(cell.text)
                    
    return "\n".join(text_parts)

def extract_cv_text(file_name: str, file_bytes: bytes) -> str:
    """
    Detect the file type and extract text from it.
    Supports PDF (.pdf) and Word (.docx) files.
    """
    file_name_lower = file_name.lower()
    
    if file_name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif file_name_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF (.pdf) or Word (.docx) file.")
